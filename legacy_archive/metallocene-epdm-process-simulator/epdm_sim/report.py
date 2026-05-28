"""Excel and PDF report export."""

from __future__ import annotations

from io import BytesIO
from typing import Any

import pandas as pd


def export_excel(result, sensitivity_df: pd.DataFrame | None = None, optimization: Any | None = None) -> bytes:
    """Build an Excel workbook containing streams, units and KPI tables."""
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        result.stream_table().to_excel(writer, sheet_name="stream table", index=False)
        result.unit_table().to_excel(writer, sheet_name="unit operations", index=False)
        pd.DataFrame([result.kpis]).drop(columns=["recommendations"], errors="ignore").to_excel(
            writer, sheet_name="product properties", index=False
        )
        result.reactor.stage_dataframe().to_excel(writer, sheet_name="reactor profile", index=False)
        result.flash1.split_table.to_excel(writer, sheet_name="flash1 split", index=False)
        result.flash2.split_table.to_excel(writer, sheet_name="flash2 split", index=False)
        result.heat_balance_table().to_excel(writer, sheet_name="heat balance", index=False)
        result.fluid_property_table().to_excel(writer, sheet_name="fluid properties", index=False)
        result.pipe_hydraulics_table().to_excel(writer, sheet_name="pressure drop", index=False)
        if sensitivity_df is not None and not sensitivity_df.empty:
            sensitivity_df.to_excel(writer, sheet_name="sensitivity", index=False)
        if optimization is not None:
            pd.DataFrame([optimization.kpis]).drop(columns=["recommendations"], errors="ignore").to_excel(
                writer, sheet_name="optimization", index=False
            )
    return buffer.getvalue()


def export_pdf_report(result, sensitivity_df: pd.DataFrame | None = None, optimization: Any | None = None) -> bytes:
    """Build a compact PDF report with assumptions, inputs and key results."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase import pdfmetrics
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    base = ParagraphStyle("cn", parent=styles["BodyText"], fontName="STSong-Light", fontSize=9, leading=13)
    title = ParagraphStyle("title", parent=styles["Title"], fontName="STSong-Light", fontSize=16)
    story = [Paragraph("茂金属乙丙橡胶 EPDM/EPM 溶液聚合工艺仿真报告", title), Spacer(1, 0.3 * cm)]
    cfg = result.config
    story.append(Paragraph("输入条件", base))
    inputs = [
        ["反应温度/°C", f"{cfg.temperature_C:.1f}", "反应压力/MPa", f"{cfg.pressure_MPa:.2f}"],
        ["乙烯 kg/h", f"{cfg.ethylene_kg_h:.2f}", "丙烯 kg/h", f"{cfg.propylene_kg_h:.2f}"],
        ["ENB kg/h", f"{cfg.enb_kg_h:.2f}", "氢气 g/h", f"{cfg.hydrogen_g_h:.2f}"],
        ["溶剂", cfg.solvent, "停留时间/min", f"{cfg.residence_time_min:.1f}"],
    ]
    story.append(_table(inputs))
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("关键结果", base))
    k = result.kpis
    outputs = [
        ["聚合物 kg/h", f"{k['polymer_kg_h']:.3f}", "固含 wt%", f"{k['solids_wt']:.2f}"],
        ["C2 wt%", f"{k['C2_wt']:.2f}", "ENB wt%", f"{k['ENB_wt']:.2f}"],
        ["Mw", f"{k['Mw']:.0f}", "PDI", f"{k['PDI']:.2f}"],
        ["门尼 ML(1+4)", f"{k['Mooney']:.1f}", "Tg/°C", f"{k['Tg_C']:.1f}"],
        ["反应热负荷/kW", f"{k['heat_duty_kW']:.2f}", "绝热温升/K", f"{k['deltaT_ad_K']:.1f}"],
        ["预热/脱挥负荷 kW", f"{k['preheat_kW']:.2f} / {k['devol_duty_kW']:.2f}", "总冷却负荷/kW", f"{k['total_cooling_load_kW']:.2f}"],
        ["黏度/Pa.s", f"{k['dynamic_viscosity_Pa_s']:.4g}", "挂胶风险", str(k["fouling_risk"])],
        ["压降/kPa", f"{k['pipe_pressure_drop_kPa']:.2f}", "泵功率/kW", f"{k['pump_power_kW']:.3f}"],
        ["ENB残留/ppm", f"{k['ENB_residue_ppm']:.0f}", "最佳牌号", f"{k['best_grade']} ({k['best_grade_score']:.1f})"],
    ]
    story.append(_table(outputs))
    story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("工艺优化建议", base))
    for rec in k["recommendations"]:
        story.append(Paragraph(f"- {rec}", base))
    if sensitivity_df is not None and not sensitivity_df.empty:
        story.append(Spacer(1, 0.25 * cm))
        story.append(Paragraph(f"敏感性分析已导出 {len(sensitivity_df)} 条计算结果。", base))
    if optimization is not None:
        story.append(Paragraph(f"优化目标牌号 {optimization.grade_id}，匹配分数 {optimization.score:.1f}。", base))
    story.append(Spacer(1, 0.25 * cm))
    story.append(
        Paragraph(
            "模型假设与局限性：本软件是研发级表观模型，不是工业设计包；动力学参数需要更多实验数据校准；"
            "热力学模型对聚合物溶液、ENB活度和高压气液平衡做了简化；流体物性默认值为工程估算，"
            "黏度、导热和密度需要实测数据补充校准；结果用于研发趋势判断、实验设计和工艺窗口筛选。",
            base,
        )
    )
    doc.build(story)
    return buffer.getvalue()


def export_word_report(result, sensitivity_df: pd.DataFrame | None = None, optimization: Any | None = None) -> bytes:
    """Build a Word report including heat balance and fluid property tables."""
    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_heading("茂金属乙丙橡胶 EPDM/EPM 溶液聚合工艺仿真报告", level=1)
    cfg = result.config
    document.add_heading("输入条件", level=2)
    _docx_table(
        document,
        [
            ["反应温度/°C", f"{cfg.temperature_C:.1f}", "反应压力/MPa", f"{cfg.pressure_MPa:.2f}"],
            ["乙烯 kg/h", f"{cfg.ethylene_kg_h:.2f}", "丙烯 kg/h", f"{cfg.propylene_kg_h:.2f}"],
            ["ENB kg/h", f"{cfg.enb_kg_h:.2f}", "氢气 g/h", f"{cfg.hydrogen_g_h:.2f}"],
            ["U W/m2/K", f"{cfg.heat_transfer_U_W_m2K:.1f}", "A m2", f"{cfg.heat_transfer_area_m2:.2f}"],
        ],
    )
    k = result.kpis
    document.add_heading("关键结果", level=2)
    _docx_table(
        document,
        [
            ["聚合物 kg/h", f"{k['polymer_kg_h']:.3f}", "固含 wt%", f"{k['solids_wt']:.2f}"],
            ["C2 wt%", f"{k['C2_wt']:.2f}", "ENB wt%", f"{k['ENB_wt']:.2f}"],
            ["反应热 kW", f"{k['heat_duty_kW']:.2f}", "绝热温升 K", f"{k['deltaT_ad_K']:.1f}"],
            ["预热/脱挥 kW", f"{k['preheat_kW']:.2f} / {k['devol_duty_kW']:.2f}", "冷却负荷 kW", f"{k['total_cooling_load_kW']:.2f}"],
            ["黏度 Pa.s", f"{k['dynamic_viscosity_Pa_s']:.4g}", "挂胶风险", str(k["fouling_risk"])],
            ["压降 kPa", f"{k['pipe_pressure_drop_kPa']:.2f}", "泵功率 kW", f"{k['pump_power_kW']:.3f}"],
        ],
    )
    document.add_heading("热量衡算表", level=2)
    _docx_dataframe(document, result.heat_balance_table())
    document.add_heading("流体性质表", level=2)
    _docx_dataframe(document, result.fluid_property_table())
    document.add_heading("压降与泵送", level=2)
    _docx_dataframe(document, result.pipe_hydraulics_table())
    document.add_heading("工艺优化建议", level=2)
    for rec in k["recommendations"]:
        document.add_paragraph(rec, style="List Bullet")
    if sensitivity_df is not None and not sensitivity_df.empty:
        document.add_paragraph(f"敏感性分析已导出 {len(sensitivity_df)} 条计算结果。")
    if optimization is not None:
        document.add_paragraph(f"优化目标牌号 {optimization.grade_id}，匹配分数 {optimization.score:.1f}。")
    document.add_heading("模型假设和数据缺口", level=2)
    document.add_paragraph(
        "本软件是研发级表观模型，不是工业设计包。反应热、流体物性、黏度模型、导热系数、压降模型均为工程估算；"
        "聚合物溶液黏度、密度、Cp、ENB活度和高压气液平衡需要实测数据进一步校准。"
    )
    document.save(buffer)
    return buffer.getvalue()


def _table(rows: list[list[str]]):
    """Create a styled reportlab table."""
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    table = Table(rows, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), "STSong-Light"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def _docx_table(document, rows: list[list[str]]) -> None:
    """Append a simple Word table from rows."""
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)


def _docx_dataframe(document, df: pd.DataFrame, max_rows: int = 20) -> None:
    """Append a compact DataFrame table to a Word document."""
    shown = df.head(max_rows)
    table = document.add_table(rows=len(shown) + 1, cols=len(shown.columns))
    table.style = "Table Grid"
    for j, col in enumerate(shown.columns):
        table.cell(0, j).text = str(col)
    for i, (_, row) in enumerate(shown.iterrows(), start=1):
        for j, col in enumerate(shown.columns):
            table.cell(i, j).text = str(row[col])
