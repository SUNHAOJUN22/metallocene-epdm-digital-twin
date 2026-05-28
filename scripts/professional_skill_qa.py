"""Executable QA harness for professional-skill-replaced peripheral workflows.

This script does not replace the scientific runtime kernel.  It codifies the
parts that are appropriate to replace with professional workflow skills:

- Excel workbook inspection that mirrors the Spreadsheets skill.
- Word report inspection that mirrors the Documents skill.
- UI contract artifact inspection that complements Browser/Playwright checks.
- Git/GitHub readiness inspection for GitHub workflow skills.
- MCP-style external tool contract inspection for future ChatGPT Apps/scientific
  simulation integrations.

The authoritative math/physics checks remain pytest, ResidualSystem,
benchmarks and release_gate.py.
"""

from __future__ import annotations

import csv
import json
import subprocess
import sys
import time
from dataclasses import asdict, dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tmp_smoke_outputs"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

REQUIRED_EXCEL_SHEETS = [
    "residual_system",
    "unit_conversion_trace",
    "benchmark_acceptance",
    "evidence_chain",
    "property_model_runtime",
    "dynamic_solver_policy",
]
FORMULA_ERROR_TOKENS = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A")


@dataclass
class SkillQaResult:
    """One professional-skill QA check result."""

    workstream: str
    professional_skill: str
    artifact: str
    passed: bool
    severity: str
    detail: str


def _latest_artifact(pattern: str) -> Path | None:
    candidates = sorted(OUT.glob(pattern), key=lambda path: path.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def _display_path(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def _run_git(args: list[str]) -> subprocess.CompletedProcess[str]:
    return subprocess.run(["git", *args], cwd=ROOT, capture_output=True, text=True)


def check_excel_report(path: Path | None = None) -> list[SkillQaResult]:
    """Check the latest Excel report with spreadsheet-style artifact rules."""
    workbook_path = path or _latest_artifact("*.xlsx")
    if workbook_path is None:
        return [
            SkillQaResult(
                "excel_report_qa",
                "spreadsheets",
                "tmp_smoke_outputs/*.xlsx",
                False,
                "P2",
                "No Excel report artifact was found for professional spreadsheet QA.",
            )
        ]

    workbook = load_workbook(workbook_path, data_only=False, read_only=True)
    sheet_names = list(workbook.sheetnames)
    long_names = [name for name in sheet_names if len(name) > 31]
    missing = [name for name in REQUIRED_EXCEL_SHEETS if name not in sheet_names]
    formula_errors: list[str] = []
    for sheet_name in sheet_names:
        worksheet = workbook[sheet_name]
        for row in worksheet.iter_rows():
            for cell in row:
                value = cell.value
                if isinstance(value, str) and any(token in value for token in FORMULA_ERROR_TOKENS):
                    formula_errors.append(f"{sheet_name}!{cell.coordinate}={value}")
                    if len(formula_errors) >= 20:
                        break
            if len(formula_errors) >= 20:
                break
        if len(formula_errors) >= 20:
            break

    passed = not long_names and not missing and not formula_errors
    detail = {
        "path": _display_path(workbook_path),
        "sheet_count": len(sheet_names),
        "long_sheet_names": long_names,
        "missing_required_sheets": missing,
        "formula_errors": formula_errors,
    }
    return [
        SkillQaResult(
            "excel_report_qa",
            "spreadsheets",
            _display_path(workbook_path),
            passed,
            "P2" if not passed else "info",
            json.dumps(detail, ensure_ascii=False),
        )
    ]


def check_word_report(path: Path | None = None) -> list[SkillQaResult]:
    """Check the latest Word report with document-style artifact rules."""
    docx_path = path or _latest_artifact("*.docx")
    if docx_path is None:
        return [
            SkillQaResult(
                "word_report_qa",
                "documents",
                "tmp_smoke_outputs/*.docx",
                False,
                "P2",
                "No Word report artifact was found for professional document QA.",
            )
        ]

    document = Document(docx_path)
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    tables = list(document.tables)
    table_texts = [
        cell.text.strip()
        for table in tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    evidence_tokens = ("residual", "risk", "benchmark", "残差", "风险", "守恒", "可信度", "审计")
    has_risk_or_residual_content = any(
        any(token in text.lower() for token in evidence_tokens) for text in [*paragraphs, *table_texts]
    )
    passed = bool(paragraphs) and bool(tables) and has_risk_or_residual_content
    detail = {
        "path": _display_path(docx_path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "table_text_cells": len(table_texts),
        "has_risk_or_residual_content": has_risk_or_residual_content,
    }
    return [
        SkillQaResult(
            "word_report_qa",
            "documents",
            _display_path(docx_path),
            passed,
            "P2" if not passed else "info",
            json.dumps(detail, ensure_ascii=False),
        )
    ]


def check_ui_contract_artifacts() -> list[SkillQaResult]:
    """Check UI smoke artifacts that complement Browser/Playwright skill QA."""
    smoke_path = OUT / "ui_e2e_smoke.json"
    workflow_path = OUT / "ui_e2e_workflow.json"
    missing = [str(path.relative_to(ROOT)) for path in [smoke_path, workflow_path] if not path.exists()]
    if missing:
        return [
            SkillQaResult(
                "ui_browser_contract_qa",
                "browser/playwright",
                ", ".join(missing),
                False,
                "P2",
                f"Missing UI contract artifacts: {', '.join(missing)}",
            )
        ]

    smoke = json.loads(smoke_path.read_text(encoding="utf-8"))
    workflow = json.loads(workflow_path.read_text(encoding="utf-8"))
    page_count = int(
        smoke.get("page_count", 0)
        or smoke.get("pages_registered", 0)
        or workflow.get("page_count", 0)
        or workflow.get("pages_registered", 0)
        or 0
    )
    manual_actions = int(workflow.get("manual_action_count", 0) or 0)
    missing_mappings = workflow.get("missing_task_mappings", []) or workflow.get("manual_actions_missing_task", []) or workflow.get("heavy_manual_without_task", [])
    heavy_export_actions = workflow.get("heavy_export_actions", []) or workflow.get("export_actions_heavy", [])
    passed = page_count > 0 and manual_actions > 0 and not missing_mappings and not heavy_export_actions
    detail = {
        "page_count": page_count,
        "manual_action_count": manual_actions,
        "missing_task_mappings": missing_mappings,
        "heavy_export_actions": heavy_export_actions,
    }
    return [
        SkillQaResult(
            "ui_browser_contract_qa",
            "browser/playwright",
            "tmp_smoke_outputs/ui_e2e_*.json",
            passed,
            "P2" if not passed else "info",
            json.dumps(detail, ensure_ascii=False),
        )
    ]


def check_github_workflow_readiness() -> list[SkillQaResult]:
    """Check local GitHub workflow readiness after publication."""
    inside = _run_git(["rev-parse", "--is-inside-work-tree"])
    remote = _run_git(["remote", "get-url", "origin"])
    branch = _run_git(["status", "--short", "--branch"])
    pushed = _run_git(["ls-remote", "--heads", "origin", "main"])

    is_repo = inside.returncode == 0 and inside.stdout.strip() == "true"
    remote_url = remote.stdout.strip()
    has_github_remote = "github.com/SUNHAOJUN22/metallocene-epdm-digital-twin" in remote_url
    has_main_remote = pushed.returncode == 0 and "refs/heads/main" in pushed.stdout
    clean_or_known = branch.returncode == 0 and "## main...origin/main" in branch.stdout
    passed = is_repo and has_github_remote and has_main_remote and clean_or_known
    detail = {
        "is_git_repo": is_repo,
        "origin": remote_url,
        "main_remote_head_present": has_main_remote,
        "branch_status": branch.stdout.strip(),
    }
    return [
        SkillQaResult(
            "github_workflow_qa",
            "github/yeet",
            remote_url or "origin",
            passed,
            "P2" if not passed else "info",
            json.dumps(detail, ensure_ascii=False),
        )
    ]


def check_mcp_interface_contract() -> list[SkillQaResult]:
    """Check MCP-style scientific tool contracts without running heavy tasks."""
    try:
        from epdm_sim.mcp import (
            call_mcp_tool,
            get_model_metadata,
            mcp_tool_registry,
            run_flowsheet_simulation,
            validate_simulation_input,
        )
    except Exception as exc:
        return [
            SkillQaResult(
                "mcp_interface_contract_qa",
                "chatgpt-apps/openai-docs",
                "epdm_sim.mcp",
                False,
                "P2",
                f"MCP interface import failed: {type(exc).__name__}: {exc}",
            )
        ]

    registry = mcp_tool_registry()
    metadata = get_model_metadata({})
    dry_run = run_flowsheet_simulation({"payload": {"temperature_C": 100.0, "pressure_MPa": 1.0}})
    invalid = validate_simulation_input({"units": {"pressure": "bar"}, "payload": {"temperature_C": 80.0}})
    unknown = call_mcp_tool("unknown_tool", {})
    passed = (
        bool(registry)
        and metadata.get("status") == "ok"
        and dry_run.get("status") == "not_run"
        and dry_run.get("heavy_task_executed") is False
        and invalid.get("status") == "rejected"
        and unknown.get("status") == "rejected"
    )
    detail = {
        "tool_count": len(registry),
        "metadata_status": metadata.get("status"),
        "dry_run_status": dry_run.get("status"),
        "dry_run_heavy_task_executed": dry_run.get("heavy_task_executed"),
        "invalid_unit_status": invalid.get("status"),
        "unknown_tool_status": unknown.get("status"),
    }
    return [
        SkillQaResult(
            "mcp_interface_contract_qa",
            "chatgpt-apps/openai-docs",
            "epdm_sim.mcp",
            passed,
            "P2" if not passed else "info",
            json.dumps(detail, ensure_ascii=False),
        )
    ]


def check_professional_skill_qa() -> list[SkillQaResult]:
    """Run all professional-skill-replaced peripheral QA checks."""
    results: list[SkillQaResult] = []
    results.extend(check_excel_report())
    results.extend(check_word_report())
    results.extend(check_ui_contract_artifacts())
    results.extend(check_github_workflow_readiness())
    results.extend(check_mcp_interface_contract())
    return results


def write_professional_skill_qa_artifacts(results: list[SkillQaResult]) -> None:
    """Write JSON/CSV professional skill QA evidence artifacts."""
    OUT.mkdir(exist_ok=True)
    payload = [asdict(result) for result in results]
    (OUT / "professional_skill_qa.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    with (OUT / "professional_skill_qa.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(payload[0].keys()) if payload else list(SkillQaResult.__annotations__))
        writer.writeheader()
        writer.writerows(payload)


def main() -> int:
    """CLI entry point for professional-skill-replaced peripheral QA."""
    started = time.perf_counter()
    results = check_professional_skill_qa()
    write_professional_skill_qa_artifacts(results)
    for result in results:
        status = "PASS" if result.passed else "FAIL"
        print(f"[{status}] {result.workstream} via {result.professional_skill}: {result.detail}")
    failed = [result for result in results if not result.passed]
    print(f"professional_skill_qa completed in {time.perf_counter() - started:.2f}s")
    return 0 if not failed else 1


if __name__ == "__main__":
    raise SystemExit(main())
