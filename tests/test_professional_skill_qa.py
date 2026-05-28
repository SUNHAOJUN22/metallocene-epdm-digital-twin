from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "scripts" / "professional_skill_qa.py"
SPEC = importlib.util.spec_from_file_location("professional_skill_qa", MODULE_PATH)
professional_skill_qa = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
sys.modules["professional_skill_qa"] = professional_skill_qa
SPEC.loader.exec_module(professional_skill_qa)


def test_excel_report_qa_rejects_long_sheet_names_and_missing_required_sheets(tmp_path: Path) -> None:
    workbook_path = tmp_path / "bad.xlsx"
    workbook = Workbook()
    workbook.active.title = "this_sheet_name_is_longer_than_31_chars"
    workbook.save(workbook_path)

    results = professional_skill_qa.check_excel_report(workbook_path)

    assert len(results) == 1
    assert results[0].professional_skill == "spreadsheets"
    assert not results[0].passed
    assert "long_sheet_names" in results[0].detail
    assert "missing_required_sheets" in results[0].detail


def test_excel_report_qa_accepts_required_sheets_without_formula_errors(tmp_path: Path) -> None:
    workbook_path = tmp_path / "good.xlsx"
    workbook = Workbook()
    workbook.active.title = professional_skill_qa.REQUIRED_EXCEL_SHEETS[0]
    for sheet_name in professional_skill_qa.REQUIRED_EXCEL_SHEETS[1:]:
        workbook.create_sheet(sheet_name)
    workbook.save(workbook_path)

    results = professional_skill_qa.check_excel_report(workbook_path)

    assert results[0].passed
    assert results[0].professional_skill == "spreadsheets"


def test_word_report_qa_requires_report_content_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "report.docx"
    document = Document()
    document.add_heading("Residual benchmark risk report", level=1)
    document.add_paragraph("Residual, benchmark and risk audit content.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "metric"
    table.cell(0, 1).text = "status"
    document.save(docx_path)

    results = professional_skill_qa.check_word_report(docx_path)

    assert results[0].passed
    assert results[0].professional_skill == "documents"
