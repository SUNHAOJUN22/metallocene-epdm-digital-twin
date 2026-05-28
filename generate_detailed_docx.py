import docx
import re

md_path = r"d:\codex\metallocene-epdm-digital-twin\README_Deep_Technical.md"
docx_path = r"d:\codex\metallocene-epdm-digital-twin\README_Deep_Technical.docx"

doc = docx.Document()
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # Simple markdown to docx mapping
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('**'):
        doc.add_paragraph(line) # Simplified bold handling
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        if line != '---':
            doc.add_paragraph(line)

doc.save(docx_path)
print(f"Successfully generated {docx_path}")
