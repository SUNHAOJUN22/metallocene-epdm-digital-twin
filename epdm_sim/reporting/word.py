"""Word report export module."""

from __future__ import annotations

from io import BytesIO
from typing import Any

import pandas as pd

from ..conservation import conservation_dataframe, run_conservation_checks
from ..engineering_checks import checks_dataframe, run_engineering_checks
from ..model_audit_report import build_model_audit_report
from ..model_confidence import build_model_confidence_card
from ..model_contracts import contracts_dataframe
from ..audit_trail import audit_trail_dataframe
from ..workflow_wizard import workflow_status
from ..cfd.grid_convergence import CFDGridConvergenceResult
from .render_helpers import _append_docx_figures, _docx_dataframe, _docx_table


def export_word_report(
    result,
    sensitivity_df: pd.DataFrame | None = None,
    optimization: Any | None = None,
    calibration: Any | None = None,
    doe_df: pd.DataFrame | None = None,
    scaleup_df: pd.DataFrame | None = None,
    figures: dict[str, Any] | None = None,
    parameter_estimation: Any | None = None,
    dynamic_semibatch_df: pd.DataFrame | None = None,
    case_comparison: pd.DataFrame | None = None,
    recycle_solver: Any | None = None,
    safety: Any | None = None,
    pareto_df: pd.DataFrame | None = None,
    uncertainty: Any | None = None,
    model_confidence: dict[str, Any] | pd.DataFrame | None = None,
    recipe_df: pd.DataFrame | None = None,
    task_log: pd.DataFrame | None = None,
    manifest: dict[str, Any] | None = None,
    calibration_loop_result: Any | None = None,
    model_audit: Any | None = None,
    posterior: Any | None = None,
    constrained_windows: pd.DataFrame | None = None,
    audit_records: Any | None = None,
    workflow_df: pd.DataFrame | None = None,
    cfd_grid_convergence: CFDGridConvergenceResult | pd.DataFrame | None = None,
) -> bytes:
    """Build a Word report including heat balance and fluid property tables."""
    from docx import Document
    from docx.shared import Inches

    buffer = BytesIO()
    document = Document()
    document.add_heading("茂金属乙丙橡胶 EPDM/EPM 溶液聚合工艺仿真报告", level=1)
    cfg = result.config
    document.add_heading("输入条件", level=2)
    _docx_table(
        document,
        [
            ["反应温度/°C", f"{cfg.temperature_C:.1f}", "反应压力/MPa", f"{cfg.pressure_MPa:.2f}"],
            ["乙烯 kg/h", f"{cfg.ethylene_kg_h:.2f}", "丙烯 kg/h", f"{cfg.propylene_kg_h:.2f}"],
            ["ENB kg/h", f"{cfg.enb_kg_h:.2f}", "氢气 g/h", f"{cfg.hydrogen_g_h:.2f}"],
            ["U W/m2/K", f"{cfg.heat_transfer_U_W_m2K:.1f}", "A m2", f"{cfg.heat_transfer_area_m2:.2f}"],
        ],
    )
    k = result.kpis
    document.add_heading("关键结果", level=2)
    _docx_table(
        document,
        [
            ["聚合物 kg/h", f"{k['polymer_kg_h']:.3f}", "固含 wt%", f"{k['solids_wt']:.2f}"],
            ["C2 wt%", f"{k['C2_wt']:.2f}", "ENB wt%", f"{k['ENB_wt']:.2f}"],
            ["反应热 kW", f"{k['heat_duty_kW']:.2f}", "绝热温升 K", f"{k['deltaT_ad_K']:.1f}"],
            ["预热/脱挥 kW", f"{k['preheat_kW']:.2f} / {k['devol_duty_kW']:.2f}", "冷却负荷 kW", f"{k['total_cooling_load_kW']:.2f}"],
            ["黏度 Pa.s", f"{k['dynamic_viscosity_Pa_s']:.4g}", "挂胶风险", str(k["fouling_risk"])],
            ["压降 kPa", f"{k['pipe_pressure_drop_kPa']:.2f}", "泵功率 kW", f"{k['pump_power_kW']:.3f}"],
        ],
    )
    document.add_heading("热量衡算表", level=2)
    _docx_dataframe(document, result.heat_balance_table())
    document.add_heading("流体性质表", level=2)
    _docx_dataframe(document, result.fluid_property_table())
    document.add_heading("压降与泵送", level=2)
    _docx_dataframe(document, result.pipe_hydraulics_table())
    document.add_heading("工程逻辑检查", level=2)
    _docx_dataframe(document, checks_dataframe(run_engineering_checks(result)), max_rows=30)
    document.add_heading("守恒闭合", level=2)
    conservation_checks = run_conservation_checks(result)
    _docx_dataframe(document, conservation_dataframe(conservation_checks), max_rows=30)
    document.add_heading("模型可信度", level=2)
    _docx_dataframe(document, build_model_confidence_card(result, conservation_results=conservation_checks).as_dataframe())
    if calibration_loop_result is not None:
        document.add_heading("实验设计与校准反馈", level=2)
        _docx_dataframe(document, calibration_loop_result.as_dataframe(), max_rows=10)
    if posterior is not None:
        document.add_heading("参数后验分布", level=2)
        _docx_dataframe(document, posterior.parameter_summary, max_rows=20)
    if constrained_windows is not None and not constrained_windows.empty:
        document.add_heading("工程约束工艺窗口", level=2)
        _docx_dataframe(document, constrained_windows, max_rows=20)
    if audit_records is not None:
        document.add_heading("审计追踪", level=2)
        _docx_dataframe(document, audit_trail_dataframe(audit_records), max_rows=20)
    if workflow_df is not None and not workflow_df.empty:
        document.add_heading("研发流程向导", level=2)
        _docx_dataframe(document, workflow_df, max_rows=20)
    if cfd_grid_convergence is not None:
        document.add_heading("CFD网格独立性验证", level=2)
        cfd_df = cfd_grid_convergence.as_dataframe() if hasattr(cfd_grid_convergence, "as_dataframe") else cfd_grid_convergence
        _docx_dataframe(document, cfd_df, max_rows=20)
    document.add_heading("模型审计详情", level=2)
    audit = model_audit or build_model_audit_report(result)
    _docx_dataframe(document, audit.as_dataframe())
    _docx_dataframe(document, audit.top_risks, max_rows=5)
    document.add_heading("模型适用范围", level=2)
    _docx_dataframe(document, contracts_dataframe(), max_rows=20)
    document.add_heading("工艺优化建议", level=2)
    for rec in k["recommendations"]:
        document.add_paragraph(rec, style="List Bullet")
    if sensitivity_df is not None and not sensitivity_df.empty:
        document.add_paragraph(f"敏感性分析结果数：{len(sensitivity_df)}")
    if optimization is not None:
        document.add_paragraph(f"优化目标：{optimization.grade_id}，匹配分数：{optimization.score:.1f}")
    if calibration is not None:
        document.add_heading("模型校准指标", level=2)
        _docx_dataframe(document, calibration.metrics_dataframe())
    if doe_df is not None and not doe_df.empty:
        document.add_heading("DOE实验建议", level=2)
        _docx_dataframe(document, doe_df, max_rows=30)
    if scaleup_df is not None and not scaleup_df.empty:
        document.add_heading("装置放大工程相似性", level=2)
        _docx_dataframe(document, scaleup_df, max_rows=20)
    recycle = recycle_solver or getattr(result, "recycle_solver", None)
    if recycle is not None:
        document.add_heading("回收系统闭合", level=2)
        document.add_paragraph(f"迭代次数：{recycle.convergence_iterations}；残差：{recycle.closure_error:.3g} kg/h。")
    if safety is not None:
        document.add_heading("热安全风险分析", level=2)
        _docx_dataframe(document, safety.as_dataframe())
    if uncertainty is not None:
        document.add_heading("模型卡与不确定性", level=2)
        _docx_dataframe(document, uncertainty.as_dataframe(), max_rows=20)
    elif model_confidence is not None:
        _docx_dataframe(document, model_confidence if isinstance(model_confidence, pd.DataFrame) else pd.DataFrame([model_confidence]))
    figure_notes = _append_docx_figures(document, figures or {}, Inches)
    if figure_notes:
        document.add_heading("静态图表说明", level=2)
        for note in figure_notes:
            document.add_paragraph(note, style="List Bullet")
    document.add_heading("声明与局限性", level=2)
    document.add_paragraph(
        "本仿真结果由 R&D 级表观模型生成，不作为工程设计依据。模型假设包含理想混合、表观动力学、线性化热力学简化。建议定期进行实验校准。"
    )
    document.save(buffer)
    return buffer.getvalue()
