import os
import ast
import docx

def get_module_info(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        source = f.read()
    
    try:
        tree = ast.parse(source)
    except SyntaxError:
        return None, []

    docstring = ast.get_docstring(tree)
    
    functions = []
    classes = []
    
    for node in tree.body:
        if isinstance(node, ast.FunctionDef) and not node.name.startswith('_'):
            func_doc = ast.get_docstring(node)
            functions.append({
                'name': node.name,
                'doc': func_doc.split('\n')[0] if func_doc else '无注释'
            })
        elif isinstance(node, ast.ClassDef) and not node.name.startswith('_'):
            class_doc = ast.get_docstring(node)
            classes.append({
                'name': node.name,
                'doc': class_doc.split('\n')[0] if class_doc else '无注释'
            })
            
    return docstring, classes, functions

base_dir = r"d:\codex\metallocene-epdm-digital-twin\epdm_sim"

output_lines = [
    "# Metallocene EPDM Digital Twin - 全量技术架构、部署与使用说明书",
    "",
    "本文档包含该软体实现的全量功能详细技术架构说明、逐步部署指南（Step by Step）以及详细的用户使用手册。",
    "",
    "---",
    "",
    "## 1. 逐步部署指南 (Deployment Step by Step)",
    "",
    "该数字孪生平台基于 Python 构建，前后端集成在 Streamlit 中，部署非常轻量。请按照以下步骤在您的计算机或服务器上运行：",
    "",
    "### 步骤 1：准备系统环境",
    "- 确保您的操作系统为 Windows、macOS 或 Linux。",
    "- 安装 **Python 3.9 或更高版本**（推荐 Python 3.10 或 3.11）。",
    "- 建议安装并使用虚拟环境工具（如 `venv` 或 `conda`）。",
    "",
    "### 步骤 2：获取工程代码",
    "将本项目源码文件夹 `metallocene-epdm-digital-twin` 下载或克隆至本地磁盘（如 `D:\\codex\\metallocene-epdm-digital-twin`）。",
    "",
    "### 步骤 3：安装依赖",
    "打开终端（命令提示符或 PowerShell），进入项目根目录：",
    "```bash",
    "cd /path/to/metallocene-epdm-digital-twin",
    "pip install -r requirements.txt",
    "```",
    "*(这将会安装所需的核心库：numpy、scipy、pandas、streamlit 等)*",
    "",
    "### 步骤 4：启动与运行系统",
    "在当前根目录下，执行以下命令以启动系统后端服务器和前端 Web 界面：",
    "```bash",
    "python -m streamlit run app.py",
    "```",
    "",
    "### 步骤 5：访问与使用",
    "启动成功后，控制台会输出网络地址。请打开浏览器并访问：",
    "**http://localhost:8501** 或者 **http://127.0.0.1:8501**",
    "此时您就可以看到平台的交互可视化控制台了。",
    "",
    "---",
    "",
    "## 2. 软件详细使用说明书 (User Manual)",
    "",
    "本软件主要用于茂金属 EPM/EPDM 溶液聚合工艺的研发模拟、敏感性分析、参数校准与热安全评估。通过 Web UI 即可完成操作：",
    "",
    "### 2.1 全局配置与模拟计算 (Sidebar Controls)",
    "- **全局快速输入栏 (左侧导航栏)**：您可以在此直接设置**核心反应条件**。如：",
    "  - 反应温度 (°C)",
    "  - 反应压力 (MPa)",
    "  - 乙烯、丙烯、ENB 和氢气进料量",
    "  - 反应器模式（支持批次、半连续、CSTR 串联等）",
    "- **运行模拟**：设置好参数后，点击“**运行快速流程模拟**”按钮。系统会在后台触发数理引擎，执行质量与能量守恒迭代计算。",
    "",
    "### 2.2 核心分析视图 (Pages Navigation)",
    "左侧栏提供了多个功能页面切换，计算完成后请切换视图进行监测与分析：",
    "- **数字孪生总览**：查看系统的全貌、核心KPI指标（如产率、聚合度、反应放热等）。",
    "- **釜式聚合工艺时序**：查看聚合物随时间或轴向深度的转化率与物性动态变化曲线。",
    "- **反应器与动力学**：监测催化剂活性、聚合动力学参数。 ",
    "- **分离脱挥与回收**：查看闪蒸罐 (Flash) 的 VLE 相平衡数据、溶剂回收效率。 ",
    "- **产品性能与美孚对标**：查看聚合物宏观物性预测，并和特定标准牌号做物理对照。",
    "- **CFD有限元可视化**：通过 3D 模型和彩色场分布查看反应器内部温度与浓度的预估分布。",
    "- **模型治理与可信度证书**：系统自动打分，评估当前的计算残差和物理可信度（非常关键，若标红说明物理量不守恒）。",
    "",
    "### 2.3 高级研发功能",
    "- **参数集与非线性估计**：输入中试实验室的真实数据点，触发底层优化器反向回归模型系数。",
    "- **报告导出**：随时在“报告导出”页面将您的运算结果生成 Excel 快照和审核报告供归档。",
    "",
    "---",
    "",
    "## 3. 技术架构与全量文件遍历 (Technical Architecture & Code Details)",
    "",
    "本数字孪生工具摒弃了传统的“黑盒经验公式”，底层全部基于“非线性方程残差感知（Residual-aware Equation-oriented）”的物理化学守恒框架进行约束求解，确保在极值工况下的预测仍具备可靠性。",
    "以下为工程底层 (`epdm_sim/`) 各子模块与其实现的核心函数与类清单：",
    ""
]

# Walk through directories and get functions/classes
for root, dirs, files in os.walk(base_dir):
    if "__pycache__" in root or ".pytest_cache" in root:
        continue
        
    for file in files:
        if file.endswith('.py') and file != "__init__.py":
            filepath = os.path.join(root, file)
            rel_path = os.path.relpath(filepath, base_dir)
            
            docstring, classes, functions = get_module_info(filepath)
            
            if not classes and not functions:
                continue
                
            output_lines.append(f"### 文件模块: `{rel_path}`")
            if docstring:
                clean_doc = docstring.split('\n')[0]
                output_lines.append(f"**技术描述**: {clean_doc}")
            
            if classes:
                output_lines.append("\n**核心技术类 (Classes)**:")
                for cls in classes:
                    output_lines.append(f"- `{cls['name']}`: {cls['doc']}")
            
            if functions:
                output_lines.append("\n**核心技术函数接口 (Functions)**:")
                for func in functions:
                    output_lines.append(f"- `{func['name']}()`: {func['doc']}")
            
            output_lines.append("\n---\n")

# Write MD
md_path = r"d:\codex\metallocene-epdm-digital-twin\README_Manual_And_Tech.md"
with open(md_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(output_lines))

# Write DOCX
docx_path = r"d:\codex\metallocene-epdm-digital-twin\README_Manual_And_Tech.docx"
doc = docx.Document()

for line in output_lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('**'):
        doc.add_paragraph(line) 
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('```'):
        continue
    else:
        if line != '---':
            doc.add_paragraph(line)

doc.save(docx_path)

print(f"Successfully generated {md_path} and {docx_path}")
